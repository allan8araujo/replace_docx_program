from tkinter import Tk, Label, Entry, Button, LabelFrame, IntVar, Checkbutton
import docx
from docx.shared import Pt

janela = Tk()
janela.title('Words replacer')
janela.geometry('450x145')
janela.resizable(width=False, height=False)

frame_1 = LabelFrame()
frame_1.pack()
janela.configure(bg='lightsteelblue2')
frame_1.configure(bg='slategray1')

textinho_0= Label(frame_1, text='Word to find', bg='slategray1')
textinho_0.grid(column=1, row=0)
textinho = Label(frame_1, text='Replace', bg='slategray1')
textinho.grid(column=2, row=0)
textinho1 = Label(frame_1, text='paragraph', bg='slategray1')
textinho1.grid(column=0, row=1)
textinho1 = Label(frame_1, text='tables', bg='slategray1')
textinho1.grid(column=0, row=2)
textinho2 = Label(frame_1, text='file name: ', bg='slategray1')
textinho2.grid(column=0, row=3)

word_entry_paragraph= Entry(frame_1)
word_entry_paragraph.grid(column=1,row=1)

word_entry_tables= Entry(frame_1)
word_entry_tables.grid(column=1,row=2)

replace_paragraph = Entry(frame_1)
replace_paragraph.grid(column=2, row=1)

replace_table = Entry(frame_1)
replace_table.grid(column=2, row=2)

doc = docx.Document('PONTOS_DEZEMBRO_2020.docx')
texto, tabelas = "", ''

paragraphs_num = (len(doc.paragraphs))
tables_num = len(doc.tables)

path_entry = Entry(frame_1)
path_entry.grid(column=1, row=3)
word= ''
var = IntVar()
checkbox = Checkbutton(frame_1, text=f'Clean this {word} word', variable=var, bg='slategray1')
checkbox.grid(column=2, row=7)

botao_data = Button(frame_1, width='25', height='1', text='Do it!', bg='slategray3',
                    command=lambda: [feriados_retirar(), paragraph_replace()])
botao_data.grid(column=1, row=7)

def paragraph_reading():
    global doc
    global texto
    for n in range(0, paragraphs_num):
        texto += str(doc.paragraphs[n].text) + '\n'


paragraph_reading()
phrase=''

def feriados_retirar():
    global texto
    global replace_to_paragraph
    global replace_to_tables
    global var
    if var.get() == 1:
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for p in cell.paragraphs:
                        if p.text[0:7] == 'FERIADO':
                            p.text = ''
                        texto += p.text
                        texto += '\n'
        print('Done')
    else:
        return print('not done')
def paragraph_replace():
    global texto
    global replace_to_paragraph
    global replace_to_tables
    global path
    global phrase
    replace_to_paragraph = replace_paragraph.get()
    replace_to_tables = replace_table.get()
    path = path_entry.get()
    path = str(path)
    word_to_find_paragraph = str(word_entry_paragraph.get())
    word_to_find_table=str(word_entry_tables.get())
    for n in range(0, paragraphs_num):
        if doc.paragraphs[n].text == "":
            continue
        elif word_to_find_paragraph in doc.paragraphs[n].text:
            doc.paragraphs[n].text=str(doc.paragraphs[n].text).replace(str(word_to_find_paragraph),str(replace_to_paragraph))
    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                for p in cell.paragraphs:
                    if word_to_find_table in p.text:
                        p.text = str(p.text).replace(str(word_to_find_table),str(replace_to_tables))
                    texto += p.text
                    texto += '\n'
    doc_gravacao = docx.Document()
    global paraObject
    paraObject = doc_gravacao.add_paragraph()
    doc.save(path + '.docx')
janela.mainloop()