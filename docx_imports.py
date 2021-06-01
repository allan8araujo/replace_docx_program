import docx
from docx.shared import Pt
from tkinter_imports import replace_paragraph, replace_table, path_entry, var, word_entry_paragraph,word_entry_tables

#here you need to put insted {YOUR ABSOLUTE PATH} the absolute path of your docx file, like 'c:/[...]/arquive.docx'
doc = docx.Document('{YOUR ABSOLUTE PATH}')
texto = ''
tabelas = ''
paragraphs_num = (len(doc.paragraphs))
tables_num = len(doc.tables)

#setting up font, styles and stuff
font = doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(10)

#this function is reading the paragraph 1 per 1 
def paragraph_reading():
    global doc
    global texto
    for n in range(0, paragraphs_num):
        texto += str(doc.paragraphs[n].text) + '\n'

# if you want to delete a word from a table into a word.docx file you can put the word insted of 'FERIADO'
def feriados_retirar():
    global texto
    global replace_to_paragraph
    global replace_to_tables
    global var
    if var.get() == 1:
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for p in cell.paragraphs:
                        #here you can replace the p.text[0:n] search list
                        if p.text[0:7] == 'FERIADO':
                            p.text = ''
                        texto += p.text
                        texto += '\n'
        print('Done')
    else:
        return print('not done')
    
#search inside tables and replace the word, the text variable is pass into the tkinter interface

def paragraph_replace():
    global texto
    global replace_to_paragraph
    global replace_to_tables
    global path
    global phrase
    replace_to_paragraph = replace_paragraph.get()
    replace_to_tables = replace_table.get()
    path = path_entry.get()
    path = str(path)
    word_to_find_paragraph = str(word_entry_paragraph.get())
    word_to_find_table=str(word_entry_tables.get())
    for n in range(0, paragraphs_num):
        if doc.paragraphs[n].text == "":
            continue
        elif word_to_find_paragraph in doc.paragraphs[n].text:
            doc.paragraphs[n].text=str(doc.paragraphs[n].text).replace(str(word_to_find_paragraph),str(replace_to_paragraph))
    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                for p in cell.paragraphs:
                    if word_to_find_table in p.text:
                        p.text = str(p.text).replace(str(word_to_find_table),str(replace_to_tables))
                    texto += p.text
                    texto += '\n'
    
    #calls the functions
    doc_gravacao = docx.Document()
    global paraObject
    paraObject = doc_gravacao.add_paragraph()
    doc.save(path + '.docx')
